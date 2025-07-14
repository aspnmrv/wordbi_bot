import config
import asyncio
import os, fnmatch
import requests
import re
import hashlib
import aiofiles
import random
import itertools
import matplotlib.pyplot as plt
import seaborn as sns

from pathlib import Path
from datetime import datetime

import numpy as np
from scipy.interpolate import interp1d
from telethon.tl.custom import Button
from telethon import TelegramClient, events, sync, functions
from telethon.tl.types import InputPeerChannel
from telethon.tl.types import SendMessageTypingAction
from PIL import Image, ImageDraw, ImageFont, ImageColor
from telethon.events import NewMessage
from telethon.tl.custom import Button
from typing import List, Literal, Tuple, Dict, Optional, Any
from collections import defaultdict
from tempfile import NamedTemporaryFile
from docx import Document

from bot.db_tools import _update_user_states, _get_current_user_step, _get_user_words, _update_user_choose_category
from bot.db import get_user_topics_db, get_user_level_db, update_user_stat_category_words_db
from bot.globals import TOPICS, WORDS, TRANSLATES
from paths import PATH_IMAGES, PATH_FONT


async def update_text_from_state_markup(
    markup: Any,
    state: List[str],
    topics: List[str],
    name: str
) -> Any:
    """Forming text for buttons depending on their state"""
    for elem in range(len(state)):
        if topics[elem] == name:
            if "✅" not in markup.rows[elem].buttons[0].text:
                markup.rows[elem].buttons[0].text += " ✅"
            else:
                markup.rows[elem].buttons[0].text = markup.rows[elem].buttons[0].text.split("✅")[0]
        else:
            markup.rows[elem].buttons[0].text = markup.rows[elem].buttons[0].text
    return markup


async def get_keyboard(text_keys: List[str]) -> List[List[Button]]:
    """"""
    keyboard = list()
    for key in range(len(text_keys)):
        keyboard.append([Button.text(text_keys[key], resize=True)])
    return keyboard


async def find_file(pattern: str, path: str) -> List[str]:
    """"""
    result = []
    for root, dirs, files in os.walk(path):
        for name in files:
            if fnmatch.fnmatch(name, pattern):
                result.append(os.path.join(root, name))
    return result


async def is_expected_steps(user_id: int, expected_steps: List[Any]) -> bool:
    """Checking if a user exists in certain steps"""

    current_step = await _get_current_user_step(user_id)

    if current_step in expected_steps:
        return True
    else:
        return False


async def create_img_card(text: str, filename: str) -> None:
    """Create a new card for a word"""

    base_img = Image.open(PATH_IMAGES / "test.png")

    print(f"=== create_img_card ===")
    print(f"Saving card for text: '{text}' to file: '{filename}'")

    with base_img.copy() as img:
        if len(text) < 15:
            size = 50
        elif 15 <= len(text) <= 20:
            size = 40
        else:
            size = 35
        font = ImageFont.truetype(str(PATH_FONT / "Unbounded-Regular.ttf"), size)
        draw = ImageDraw.Draw(img)
        w, h = 780, 502
        draw.text((w // 2, h // 2), text, font=font, fill="black", anchor="mm")
        img.save(PATH_IMAGES / filename)


async def build_img_cards(words: Dict[str, str]) -> None:
    """"""

    for word, word_ru in words.items():
        await create_img_card(word.lower(), f"{word.replace(' ', '').lower()}_en.png")
        await create_img_card(word_ru.lower(), f"{word.replace(' ', '').lower()}_ru.png")

    return


async def get_proposal_topics(topics: List[str], states: Optional[List[str]] = None) -> List[List[Button]]:
    """"""
    if states is None:
        buttons = [[Button.inline(text=topic, data=topic)] for topic in topics]
    else:
        buttons = [[Button.inline(text=topic + " " + state, data=topic)] for topic, state in zip(topics, states)]

    return buttons


async def build_markup(topics: List[str], current_state: List[str]) -> List[List[Button]]:
    """"""
    markup = [
        [
            Button.inline(text=topics[i] + current_state[i], data=topics[i])
        ] for i in range(len(topics))
    ]

    return markup


async def get_state_markup(markup: Any, user_id: int) -> List[str]:
    """Forming buttons and their states for the user"""

    state = list()

    for row in range(len(markup.rows)):
        text = markup.rows[row].buttons[0].text
        state.append('✅' if '✅' in text else "")

    await _update_user_states(user_id, "states", state)

    return state


async def match_topics_name(topics: List[str]) -> List[str]:
    """Topic name matching"""

    match_topics = list()
    for topic in topics:
        match_topics.append(TOPICS[topic])
    return match_topics


async def get_diff_between_ts(last_ts: Optional[str]) -> float:
    """Returns the difference between two timestamps in seconds"""

    if last_ts is not None:
        current_time = datetime.now()
        last_ts = datetime.strptime(last_ts, "%Y-%m-%d %H:%M:%S.%f")

        return (current_time - last_ts).total_seconds()
    else:
        return 1000


async def build_list_of_words(user_topics: List[str], user_level: str, user_id: int = None):
    """Building a word list for the user depending on the level"""

    total_words = list()
    word_to_category = dict()

    if "Beginner" in user_level:
        level_key = "a"
    elif "Intermediate" in user_level:
        level_key = "b"
    else:
        level_key = "c"

    for topic, values in WORDS.items():
        if topic in user_topics:
            words = values[level_key]
            total_words.append(words)
            for word in words:
                word_to_category[word] = f"{TOPICS[topic]}_{level_key.upper()}_level"

    print("word_to_category", word_to_category)

    bulk_data = {}
    for word, category in word_to_category.items():
        if category not in bulk_data:
            bulk_data[category] = {}
        bulk_data[category][word] = TRANSLATES[word]

    for category, words_dict in bulk_data.items():
        await update_user_stat_category_words_db(
            user_id=user_id,
            words=words_dict,
            category=category,
            is_system=True
        )

    return word_to_category


async def build_history_message(data: List[Tuple[str, str, str]]) -> List[Dict[str, str]]:
    """"""
    result = list()
    for b in data:
        if b[0] == "user":
            item = {"role": "user", "content": b[2]}
        else:
            item = {"role": "assistant", "content": b[2]}
        result.append(item)

    return result


async def check_exist_img(file_name: str) -> bool:
    """Checking if a file with this name exists"""

    file_path = os.path.join("images", file_name)

    return os.path.isfile(file_path)


async def send_img(
    event: NewMessage.Event,
    buttons: List[List[Button]],
    file_name: str,
    current_word: str,
    lang: str,
    type_action: Literal["send", "edit"]
) -> None:
    """Sending an image"""

    if await check_exist_img(file_name):
        if type_action == "send":
            await event.client.send_message(event.chat_id, buttons=buttons,
                                            file=f"/{PATH_IMAGES}/{file_name}")
        elif type_action == "edit":
            await event.client.edit_message(event.sender_id, event.original_update.msg_id,
                                            file=f"/{PATH_IMAGES}/{file_name}",
                                            buttons=buttons)
        else:
            raise ValueError(f"Invalid action type: {type_action}. Expected 'send' or 'edit'.")
    else:
        await create_img_card(current_word.lower(), file_name)
        if type_action == "send":
            await event.client.send_message(event.chat_id, buttons=buttons,
                                            file=f"/{PATH_IMAGES}/{file_name}")
        elif type_action == "edit":
            await event.client.edit_message(event.sender_id, event.original_update.msg_id,
                                            file=f"/{PATH_IMAGES}/{file_name}",
                                            buttons=buttons)
        else:
            raise ValueError(f"Invalid action type: {type_action}. Expected 'send' or 'edit'.")
    return


async def get_translate_word(word: str, from_lang: str) -> str:
    """Translation of a word from a dictionary"""

    if from_lang == "en":
        return TRANSLATES[word]
    else:
        raise ValueError(f"It is impossible to translate into this language")


async def get_code_fill_form(user_id: int) -> int:
    """Returns a status code depending on the availability of data"""

    user_topics = await get_user_topics_db(user_id)
    user_level = await get_user_level_db(user_id)
    user_words = await _get_user_words(user_id)
    if not user_topics and not user_level:
        return -1
    elif not user_topics:
        return -2
    elif not user_level:
        return -3
    elif not user_words:
        return -4
    return 0


async def get_users_info(user_id: int):
    """"""

    user_topics = await get_user_topics_db(user_id)
    user_level = await get_user_level_db(user_id)
    user_words = await _get_user_words(user_id)

    if not user_topics and not user_level:
        text_level = f"**Уровень:** пока не выбран\n\n"
        text_topics = "**Выбранные интересы:** пока не выбраны\n\n"
    else:
        if user_level:
            text_level = f"**Уровень:** {user_level}\n\n"
        else:
            text_level = f"**Уровень:** пока не выбран\n\n"
        if user_topics:
            text_topics = "**Выбранные интересы:**\n\n" + "\n".join(f"▪️ {interest}" for interest in user_topics) + "\n\n"
        else:
            text_topics = "**Выбранные интересы:** пока не выбраны\n\n"

    text = "⚙️ Текущие настройки:\n\n" + text_level + text_topics

    return text


async def send_user_file_stat(event, file, text):
    """"""
    if file:
        await event.client.send_file(
            event.chat_id,
            file.name,
            caption=text,
            silent=True
        )

    return


async def draw_words_line_chart(data):
    """"""
    if len(data) <= 1:
        return

    dates_raw = [d[0] for d in data]
    counts = [d[1] for d in data]
    x = np.arange(len(dates_raw))

    if len(x) >= 4:
        x_smooth = np.linspace(x.min(), x.max(), 300)
        f = interp1d(x, counts, kind="cubic")
        y_smooth = f(x_smooth)
        smooth_x = x_smooth
        smooth_y = y_smooth
    else:
        smooth_x = x
        smooth_y = counts

    date_labels = [d.strftime("%b %d") for d in dates_raw]

    sns.set_theme(style="white")
    fig, ax = plt.subplots(figsize=(10, 5), dpi=180)
    ax.plot(smooth_x, smooth_y, color="#b497bd", linewidth=2)
    ax.scatter(x, counts, color="#b497bd", s=60, zorder=5)

    ax.set_facecolor("white")
    fig.patch.set_facecolor("white")
    for spine in ["top", "right", "bottom"]:
        ax.spines[spine].set_visible(False)

    ax.spines["left"].set_linewidth(0.8)
    ax.spines["left"].set_color("#aaa")
    ax.grid(False)

    ax.set_xticks(x)
    ax.set_xticklabels(date_labels, rotation=45, fontsize=10)
    ax.set_ylabel("Words", fontsize=11)
    ax.tick_params(axis="x", bottom=False, labelbottom=True)
    ax.tick_params(axis="y", left=True, labelleft=True)

    plt.tight_layout()
    file = NamedTemporaryFile(delete=False, suffix=".png")
    plt.savefig(file.name)
    plt.clf()
    plt.close()

    return file


async def draw_words_category_chart(data):
    """"""
    if len(data) <= 1:
        return

    categories = [row[0].lower() for row in data if row[0]]
    words = [row[1] for row in data if row[0]]
    learned_words = [row[2] for row in data if row[0]]
    shares = [row[3] for row in data if row[0]]

    sns.set_theme(style="white")
    fig, ax = plt.subplots(figsize=(12, 6), dpi=180)

    bar_width = 0.6
    x = range(len(categories))

    for i, (cat, total, learned, share) in enumerate(zip(categories, words, learned_words, shares)):
        ax.bar(i, total, width=bar_width, color="#d8cbe6", edgecolor="none")
        ax.bar(i, total * share / 100, width=bar_width, color="#b497bd", edgecolor="none")
        ax.text(i, total + 0.05, f"{total}", ha="center", va="bottom", fontsize=12, fontweight="bold")
        ax.text(i, (total * share / 100) * 0.65, f"{int(share)}%",
                ha="center", va="center", fontsize=10, color="white", fontweight="medium")

    ax.set_facecolor("white")
    fig.patch.set_facecolor("white")
    for spine in ax.spines.values():
        spine.set_visible(False)
    ax.tick_params(left=False, bottom=False, labelleft=False)

    ax.grid(False)
    ax.set_xticks(x)
    ax.set_xticklabels(categories, fontsize=12)

    plt.tight_layout()
    file = NamedTemporaryFile(delete=False, suffix=".png")
    plt.savefig(file.name)
    plt.clf()
    plt.close()

    return file


# async def extract_text_from_docx(path):
#     """"""
#     doc = Document(path)
#     full_text = []
#     for para in doc.paragraphs:
#         full_text.append(para.text)
#     return '\n'.join(full_text)


async def extract_text_from_docx(path):
    ext = os.path.splitext(path)[1].lower()

    if ext == ".txt":
        async with aiofiles.open(path, mode="r", encoding="utf-8") as f:
            content = await f.read()
        return content

    elif ext == ".docx":
        loop = asyncio.get_running_loop()
        return await loop.run_in_executor(None, lambda: "\n".join(
            para.text for para in Document(path).paragraphs
        ))

    else:
        raise ValueError(f"Unsupported file type: {ext}")


async def is_valid_word_list(text):
    """"""
    lines = text.strip().split('\n')
    for line in lines:
        if ":" not in line:
            return False
        parts = line.split(":", 1)
        if len(parts) != 2 or not parts[0].strip() or not parts[1].strip():
            return False
    return True


async def parse_word_list(text):
    """"""
    word_dict = {}
    lines = text.strip().split('\n')
    for line in lines:
        key, value = line.split(':', 1)
        word_dict[key.strip()] = value.strip()
    return word_dict


async def is_simple_word_list(text):
    """
    """
    lines = [line.strip() for line in text.strip().split('\n') if line.strip()]
    if len(lines) < 2:
        return False
    for line in lines:
        if ':' in line:
            return False
    return True


async def cut_word_pairs(content, limit=35):
    """"""
    lines = [line.strip() for line in content.splitlines() if line.strip()]
    return "\n".join(lines[:limit])


async def generate_uid(user_id, categories):
    """
    user_id: int или str
    categories: list[str], e.g. ['sport', 'books']
    """
    sorted_cats = sorted(categories)
    combined_str = f"{user_id}|" + "|".join(sorted_cats)

    uid_hash = hashlib.sha256(combined_str.encode()).hexdigest()

    return uid_hash[:5]


async def is_valid_base_date(word):
    """"""
    pattern = r"^base_\d{4}-\d{2}-\d{2}.*$"

    return bool(re.match(pattern, word))


async def normalize_filename(word):
    return word.strip().replace(' ', '').lower()
